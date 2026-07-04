from pathlib import Path

import pytest
from docx import Document as DocxDocument

from app import document_loader


def test_load_document_file_returns_text_and_metadata(tmp_path):
    raw_file = tmp_path / "sample.txt"
    raw_file.write_text("  alpha\t\tbeta  ", encoding="utf-8")

    document = document_loader.load_document_file(raw_file, base_path=tmp_path)

    assert document == {
        "text": "alpha beta",
        "metadata": {
            "source": "sample.txt",
            "filename": "sample.txt",
            "relative_path": "sample.txt",
        },
    }


def test_load_document_file_loads_markdown_as_plain_text(tmp_path):
    raw_file = tmp_path / "notes.md"
    raw_file.write_text("# Title\n\nRAG uses **context**.", encoding="utf-8")

    document = document_loader.load_document_file(raw_file, base_path=tmp_path)

    assert document["text"] == "# Title\n\nRAG uses **context**."
    assert document["metadata"]["filename"] == "notes.md"
    assert document["metadata"]["relative_path"] == "notes.md"


def test_load_document_file_raises_for_missing_file(tmp_path):
    missing_file = tmp_path / "missing.txt"

    with pytest.raises(FileNotFoundError, match="document file does not exist"):
        document_loader.load_document_file(missing_file, base_path=tmp_path)


def test_load_document_file_raises_for_directory(tmp_path):
    with pytest.raises(ValueError, match="document path must be a file"):
        document_loader.load_document_file(tmp_path, base_path=tmp_path)


def test_load_document_file_raises_for_unsupported_extension(tmp_path):
    raw_file = tmp_path / "ignored.json"
    raw_file.write_text("{}", encoding="utf-8")

    with pytest.raises(ValueError, match="unsupported document extension"):
        document_loader.load_document_file(raw_file, base_path=tmp_path)


def test_load_document_file_raises_for_empty_text(tmp_path):
    raw_file = tmp_path / "empty.md"
    raw_file.write_text(" \t\n ", encoding="utf-8")

    with pytest.raises(ValueError, match="document is empty"):
        document_loader.load_document_file(raw_file, base_path=tmp_path)


def test_load_document_file_raises_clear_unicode_error(tmp_path):
    raw_file = tmp_path / "bad.txt"
    raw_file.write_bytes(b"\xff\xfe\x00")

    with pytest.raises(ValueError, match="failed to decode document as UTF-8"):
        document_loader.load_document_file(raw_file, base_path=tmp_path)


def test_txt_collapses_spaces_and_tabs_but_preserves_newlines(tmp_path):
    raw_file = tmp_path / "sample.txt"
    raw_file.write_text("  alpha\t\tbeta  \nsecond\t line  ", encoding="utf-8")

    documents = document_loader.load_documents(str(tmp_path))

    assert documents == [
        {
            "text": "alpha beta \nsecond line",
            "source": "sample.txt",
        }
    ]


def test_empty_txt_file_is_skipped(tmp_path):
    raw_file = tmp_path / "empty.txt"
    raw_file.write_text(" \t  ", encoding="utf-8")

    documents = document_loader.load_documents(str(tmp_path))

    assert documents == []


def test_markdown_files_are_loaded_by_load_documents(tmp_path):
    raw_file = tmp_path / "notes.md"
    raw_file.write_text("markdown note", encoding="utf-8")

    documents = document_loader.load_documents(str(tmp_path))

    assert documents == [
        {
            "text": "markdown note",
            "source": "notes.md",
        }
    ]


def test_docx_paragraphs_are_joined_with_newlines(tmp_path):
    raw_file = tmp_path / "notes.docx"
    doc = DocxDocument()
    doc.add_paragraph("First   paragraph")
    doc.add_paragraph("Second\tparagraph")
    doc.save(raw_file)

    documents = document_loader.load_documents(str(tmp_path))

    assert documents == [
        {
            "text": "First paragraph\nSecond paragraph",
            "source": "notes.docx",
        }
    ]


def test_pdf_returns_one_document_and_skips_empty_pages(tmp_path, monkeypatch):
    raw_file = tmp_path / "paper.pdf"
    raw_file.write_bytes(b"fake pdf content")

    class FakePage:
        def __init__(self, text):
            self.text = text

        def extract_text(self):
            return self.text

    class FakePdfReader:
        def __init__(self, file_path):
            assert Path(file_path) == raw_file
            self.pages = [
                FakePage(" First   page "),
                FakePage(None),
                FakePage("   "),
                FakePage("Second\tpage"),
            ]

    monkeypatch.setattr(document_loader, "PdfReader", FakePdfReader, raising=False)

    documents = document_loader.load_documents(str(tmp_path))

    assert documents == [
        {
            "text": "First page\nSecond page",
            "source": "paper.pdf",
        }
    ]


def test_recursive_scan_uses_posix_relative_source_paths(tmp_path):
    first = tmp_path / "first" / "same.txt"
    second = tmp_path / "second" / "same.txt"
    first.parent.mkdir()
    second.parent.mkdir()
    first.write_text("first file", encoding="utf-8")
    second.write_text("second file", encoding="utf-8")

    documents = document_loader.load_documents(str(tmp_path))

    assert {"text": "first file", "source": "first/same.txt"} in documents
    assert {"text": "second file", "source": "second/same.txt"} in documents


def test_unsupported_files_are_ignored(tmp_path):
    raw_file = tmp_path / "ignored.json"
    raw_file.write_text("ignored", encoding="utf-8")

    documents = document_loader.load_documents(str(tmp_path))

    assert documents == []


def test_failed_file_load_prints_warning_and_is_skipped(tmp_path, capsys):
    raw_file = tmp_path / "broken.pdf"
    raw_file.write_bytes(b"not a valid pdf")

    documents = document_loader.load_documents(str(tmp_path))

    captured = capsys.readouterr()
    assert documents == []
    assert "[Loader Warning] Failed to load" in captured.out
    assert "broken.pdf" in captured.out


def test_missing_folder_returns_empty_list(tmp_path):
    missing_folder = tmp_path / "missing"

    documents = document_loader.load_documents(str(missing_folder))

    assert documents == []
